from docx import Document
from docx.shared import Pt
import os
from src.utils.log import logger

def read_word_file(file):
    document = Document(file)
    
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)


def extract_text_from_page(pdf_reader,
                           page_num:int) -> str:
    page = pdf_reader.pages[page_num]
    return page.extract_text()


def split_docx(file_path):
    # Load the document
    doc = Document(file_path)
    
    # Create a directory to save the split pages
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output_dir = f"{base_name}_pages"
    os.makedirs(output_dir, exist_ok=True)

    # Initialize a new Document for splitting
    page_doc = Document()
    
    # Iterate through paragraphs and add them to the new document
    page_number = 1
    for paragraph in doc.paragraphs:
        page_doc.add_paragraph(paragraph.text)
        
        # Check if the paragraph is a page break
        if paragraph.text == "":
            # Save the current page
            page_file_name = os.path.join(output_dir, f"{base_name}_page_{page_number}.docx")
            page_doc.save(page_file_name)
            print(f"Saved: {page_file_name}")
            
            # Start a new document for the next page
            page_doc = Document()
            page_number += 1
    
    # Save any remaining content as the last page
    if page_doc.paragraphs:
        page_file_name = os.path.join(output_dir, f"{base_name}_page_{page_number}.docx")
        page_doc.save(page_file_name)
        print(f"Saved: {page_file_name}")


def extract_images_from_doc(pdf_path):
    images = convert_from_path(pdf_path)
    image_files = []
    for i, image in enumerate(images):
        image_file = f'image_{i + 1}.png'
        image.save(image_file, 'PNG')
        image_files.append(image_file)
    return image_files





def save_translated_docx(translated_pages,
                         original_pdf):
    
    output_pdf = io.BytesIO()
    pdf_writer = PyPDF2.PdfWriter()
    
    for page_num, translated_text in sorted(translated_pages.items()):
        page = original_pdf.pages[page_num]
        pdf_writer.add_blank_page(width=page.mediabox.width,
                                height=page.mediabox.width)
        # Ajouter le texte traduit à la page PDF
        pdf_writer.add_page 
        # pages[page_num].merge_page(translated_text, expand = True)

    pdf_writer.write(output_pdf)
    output_pdf.seek(0)
    return output_pdf
